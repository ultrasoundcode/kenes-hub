from datetime import datetime
from decimal import Decimal
from django.template import Template, Context
from django.conf import settings
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
import os
from .models import GeneratedDocument, DocumentTemplate
from apps.applications.models import Application
from apps.accounts.models import User


class DocumentService:
    """Сервис для генерации документов"""

    @staticmethod
    def generate_document(template_id, application_id, user_id, field_values=None):
        """Генерирует документ на основе шаблона"""
        
        template = DocumentTemplate.objects.get(id=template_id)
        application = Application.objects.get(id=application_id)
        user = User.objects.get(id=user_id)
        
        # Подготавливаем данные для подстановки
        context_data = DocumentService._prepare_context(application, field_values or {})
        
        # Определяем тип шаблона и генерируем документ
        if template.template_file.name.endswith('.docx'):
            return DocumentService._generate_from_docx(template, context_data, application, user)
        else:
            return DocumentService._generate_from_text(template, context_data, application, user)

    @staticmethod
    def _prepare_context(application, field_values):
        """Подготавливает контекст для шаблона"""
        
        applicant = application.applicant
        creditor = application.creditor_name
        
        # Базовые поля
        context = {
            # Заявка
            'application_number': application.number,
            'application_subject': application.subject,
            'application_description': application.description,
            'application_created': application.created_at.strftime('%d.%m.%Y'),
            'application_amount': application.amount,
            'creditor_name': creditor or '',
            'contract_number': application.contract_number or '',
            
            # Заявитель
            'applicant_fio': f"{applicant.last_name} {applicant.first_name}",
            'applicant_last_name': applicant.last_name,
            'applicant_first_name': applicant.first_name,
            'applicant_middle_name': getattr(applicant, 'middle_name', ''),
            'applicant_iin': applicant.iin or '',
            'applicant_phone': applicant.phone or '',
            'applicant_email': applicant.email,
            'applicant_address': getattr(applicant.profile, 'address', '') if hasattr(applicant, 'profile') else '',
            
            # Текущая дата
            'current_date': datetime.now().strftime('%d.%m.%Y'),
            'current_date_long': datetime.now().strftime('%d %B %Y года'),
            'current_year': datetime.now().year,
            
            # Дополнительные данные из field_values
            **field_values
        }
        
        return context

    @staticmethod
    def _generate_from_docx(template, context_data, application, user):
        """Генерирует документ из DOCX шаблона"""
        
        # Открываем шаблон
        doc = Document(template.template_file.path)
        
        # Заменяем плейсхолдеры в параграфах
        for paragraph in doc.paragraphs:
            for key, value in context_data.items():
                placeholder = f'{{{{ {key} }}}}'
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value or ''))
        
        # Заменяем плейсхолдеры в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in context_data.items():
                        placeholder = f'{{{{ {key} }}}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value or ''))
        
        # Сохраняем файл
        filename = f"document_{application.number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(settings.MEDIA_ROOT, 'documents', 'original', filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)
        
        # Создаем запись в БД
        generated_doc = GeneratedDocument.objects.create(
            template=template,
            application=application,
            created_by=user,
            original_file=f"documents/original/{filename}",
            document_data=context_data,
            field_values=context_data,
            status='generated'
        )
        
        return generated_doc

    @staticmethod
    def _generate_from_text(template, context_data, application, user):
        """Генерирует документ из текстового шаблона"""
        
        # Читаем содержимое шаблона
        with open(template.template_file.path, 'r', encoding='utf-8') as f:
            template_content = f.read()
        
        # Применяем шаблон Django
        template_obj = Template(template_content)
        context = Context(context_data)
        rendered_content = template_obj.render(context)
        
        # Создаем PDF
        filename = f"document_{application.number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(settings.MEDIA_ROOT, 'documents', 'original', filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        c = canvas.Canvas(filepath, pagesize=A4)
        width, height = A4
        
        # Добавляем текст
        text_object = c.beginText(2*cm, height - 2*cm)
        text_object.setFont("Helvetica", 12)
        
        for line in rendered_content.split('\n'):
            text_object.textLine(line)
        
        c.drawText(text_object)
        c.save()
        
        # Создаем запись в БД
        generated_doc = GeneratedDocument.objects.create(
            template=template,
            application=application,
            created_by=user,
            original_file=f"documents/original/{filename}",
            document_data=context_data,
            field_values=context_data,
            status='generated'
        )
        
        return generated_doc

    @staticmethod
    def clone_document(document_id, user_id):
        """Создает копию документа"""
        
        original_doc = GeneratedDocument.objects.get(id=document_id)
        user = User.objects.get(id=user_id)
        
        # Создаем новый документ на основе существующего
        new_doc = GeneratedDocument.objects.create(
            template=original_doc.template,
            application=original_doc.application,
            created_by=user,
            document_data=original_doc.document_data,
            field_values=original_doc.field_values,
            version=original_doc.version + 1,
            status='draft'
        )
        
        # Копируем файл
        if original_doc.original_file:
            import shutil
            from django.core.files.storage import default_storage
            
            old_path = original_doc.original_file.path
            new_filename = f"document_{new_doc.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            new_path = os.path.join(settings.MEDIA_ROOT, 'documents', 'original', new_filename)
            
            os.makedirs(os.path.dirname(new_path), exist_ok=True)
            shutil.copy2(old_path, new_path)
            
            new_doc.original_file = f"documents/original/{new_filename}"
            new_doc.save()
        
        return new_doc

    @staticmethod
    def get_document_context(document_id):
        """Возвращает контекст для документа"""
        
        document = GeneratedDocument.objects.get(id=document_id)
        return document.document_data

    @staticmethod
    def update_document_fields(document_id, field_values):
        """Обновляет поля документа"""
        
        document = GeneratedDocument.objects.get(id=document_id)
        document.field_values.update(field_values)
        document.save()
        
        return document